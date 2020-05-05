# -*- coding: utf-8 -*-

import random
import string
import os

from attr import attrs, attrib, fields, asdict
from datetime import datetime
from flask_pymongo import ObjectId
from typing import List, Any, Dict, Tuple
from flask import abort, current_app
from flask_login import UserMixin
from werkzeug.security import check_password_hash, generate_password_hash
from docx import Document
from transliterate import translit

from app import login, pymongo

__all__ = ["RegistrationData", "PrimaryData", "SecondaryBiomarkers", "SeriesData", "Series",
           "PatientCollection", "Patient", "User", "UserCollection"]


@attrs
class _PatientData:
    FIELD_NAME = ""

    id = None

    def is_full_filled(self) -> bool:
        return all(value is not None for value in asdict(self).values())

    def serialize(self) -> Dict[str, Any]:
        return {self.__class__.FIELD_NAME: asdict(self, filter=lambda _, value: value is not None)}

    @classmethod
    def create_from_dict(cls, data: Dict[str, Any]) -> "_PatientData":
        instance = cls(**data.get(cls.FIELD_NAME, {}))
        instance.id = data.get("_id")
        return instance


@attrs(repr=False)
class RegistrationData(_PatientData):
    FIELD_NAME = "registration_data"

    name = attrib(type=str, converter=str.title)
    surname = attrib(type=str, converter=str.title)
    birthday = attrib(type=datetime)
    mobile_number = attrib(type=str)

    @property
    def age(self) -> int:
        return int((datetime.now() - self.birthday).days / 365.25)

    def __repr__(self) -> str:
        return f"Фамилия: {self.surname}\n" \
               f"Имя: {self.name}\n" \
               f"Возраст: {self.age} лет\n" \
               f"Номер мобильного телефона: +7{self.mobile_number}"


@attrs(repr=False)
class PrimaryData(_PatientData):
    FIELD_NAME = "primary_data"

    height = attrib(type=int, default=None)
    weight = attrib(type=int, default=None)
    is_smoking = attrib(type=bool, default=None)
    complaints = attrib(type=str, default=None)

    def __repr__(self) -> str:
        return f"Рост: {self.height} см\n" \
               f"Вес: {self.weight} кг\n" \
               f"Курит?: {'да' if self.is_smoking else 'нет'}\n" \
               f"Жалобы: {self.complaints}"


@attrs(repr=False)
class SecondaryBiomarkers(_PatientData):
    FIELD_NAME = "secondary_biomarkers"

    mmse = attrib(type=int, default=None)
    moca = attrib(type=int, default=None)

    def __repr__(self) -> str:
        return f"MMSE: {self.mmse}\n" \
               f"MoCA: {self.moca}"


@attrs(repr=False)
class Series:
    id = None

    desc = attrib(type=str)
    dt = attrib(type=datetime)

    slice_count = attrib(type=int)

    dicom_path = attrib(type=str)
    nifti_dir = attrib(type=str)
    img_dir = attrib(type=str)

    # данные, полученные после проведения морфометического анализа
    whole_brain_volume = attrib(type=float, default=None)
    left_volume = attrib(type=float, default=None)
    right_volume = attrib(type=float, default=None)
    status = attrib(type=str, default=None)  # ok, timeout, runtime error

    def is_full_filled(self) -> bool:
        return all(value is not None for value in asdict(self).values())

    @property
    def normed_left_volume(self) -> float:
        if self.left_volume is not None and self.whole_brain_volume is not None:
            return round(self.left_volume / self.whole_brain_volume, 5)

    @property
    def normed_right_volume(self) -> float:
        if self.right_volume is not None and self.whole_brain_volume is not None:
            return round(self.right_volume / self.whole_brain_volume, 5)

    @property
    def image_paths(self) -> List[str]:
        img_dir = os.path.join(current_app.static_folder, self.img_dir)
        sorted_list_dir = sorted(os.listdir(img_dir), key=lambda x: int(x.split('.')[0]))
        return [os.path.join(self.img_dir, basename) for basename in sorted_list_dir]

    def __repr__(self) -> str:
        return f"Дата и время создания: {self.dt}\n" \
               f"Объем левого гиппокампа: {self.left_volume} мм\u00b3\n" \
               f"Объем правого гиппокампа: {self.right_volume} мм\u00b3\n" \
               f"Объем мозговой ткани: {self.whole_brain_volume} мм\u00b3\n" \
               f"Нормализованный объем левого гиппокампа: {self.normed_left_volume}\n" \
               f"Нормализованный объем правого гиппокампа: {self.normed_right_volume}"


@attrs
class SeriesData(_PatientData):
    FIELD_NAME = "series_data"

    series_dict = attrib(type=Dict[str, Series], default={})

    def __len__(self) -> int:
        return len(self.series_dict)

    def is_full_filled(self) -> bool:
        return all(series.is_full_filled() for series in self.find_all())

    def find_or_404(self, series_id: str) -> Series:
        series = self.series_dict.get(series_id)
        if series is None:
            abort(404)

        instance = Series(**series)
        instance.id = series_id
        return instance

    def find_all(self) -> List[Series]:
        instances = []
        for series_id, series_data in self.series_dict.items():
            instance = Series(**series_data)
            instance.id = series_id
            instances.append(instance)

        return instances

    def insert(self, series: Series, series_id) -> None:
        self.series_dict[series_id] = series

    def remove(self, series_id: str) -> None:
        if series_id in self.series_dict:
            del self.series_dict[series_id]


@attrs
class Patient(_PatientData):
    registration_data = attrib(type=RegistrationData)
    primary_data = attrib(type=PrimaryData, default=None)
    secondary_biomarkers = attrib(type=SecondaryBiomarkers, default=None)
    series_data = attrib(type=SeriesData, default=None)

    def get_report(self) -> Document:
        assert self.registration_data.is_full_filled(), "Не все регистрационные данные заполнены"
        assert self.primary_data.is_full_filled(), "Не все первичные данные заполнены"
        assert self.secondary_biomarkers.is_full_filled(), "Не все другие биомаркеры заполнены"
        assert len(self.series_data) > 0, "Должна быть загружена хотя бы одна серия"
        assert self.series_data.is_full_filled(), "Не все загруженные серии проанализированы"

        document = Document()
        document.add_heading("Отчет", level=0)

        document.add_heading("Регистрационные данные", level=1)
        document.add_paragraph(repr(self.registration_data))

        document.add_heading("Первичные данные", level=1)
        document.add_paragraph(repr(self.primary_data))

        document.add_heading("Другие биомаркеры", level=1)
        document.add_paragraph(repr(self.secondary_biomarkers))

        document.add_heading("Серии", level=1)
        for series in self.series_data.find_all():
            document.add_heading(f"Серия: {series.desc}", level=2)
            document.add_paragraph(repr(series))

        return document

    @classmethod
    def create_from_dict(cls, data: Dict[str, Any]) -> "Patient":
        instance = cls(**{field.name: field.type(**data.get(field.type.FIELD_NAME, {})) for field in fields(Patient)})
        instance.id = data.get("_id")
        return instance


@attrs
class User(UserMixin):
    password = None

    id = attrib(type=str)
    email = attrib(type=str)
    name = attrib(type=str, converter=str.title)
    surname = attrib(type=str, converter=str.title)
    is_admin = attrib(type=bool, default=False)
    password_hash = attrib(type=str, default=None)

    def check_password(self, password: str) -> bool:
        return check_password_hash(self.password_hash, password)

    def set_new_password(self) -> None:
        password_len = random.randint(8, 16)
        self.password = ''.join(random.sample(string.ascii_letters + string.digits, k=password_len))
        self.password_hash = generate_password_hash(self.password)

    @classmethod
    def create_from_dict(cls, data: Dict[str, Any]) -> "User":
        data["id"] = data["_id"]
        del data["_id"]
        return cls(**data)

    @classmethod
    def register(cls, name: str, surname: str, email: str) -> "User":
        name_translit = translit(name, "ru", reversed=True).lower()
        surname_translit = translit(surname, "ru", reversed=True).lower()
        login = name_translit + surname_translit + str(UserCollection.docs_count())

        instance = User(id=login, email=email, name=name, surname=surname)
        instance.set_new_password()
        return instance


class UserCollection:

    @staticmethod
    def init(app) -> None:
        if "email_" not in pymongo.db.users.index_information():
            pymongo.db.users.create_index("email", name="email_", unique=True)

        admin_name = app.config["ADMIN_NAME"]
        admin_surname = app.config["ADMIN_SURNAME"]
        admin_email = app.config["ADMIN_EMAIL"]
        admin_login = app.config["ADMIN_LOGIN"]
        admin_password_hash = generate_password_hash(app.config["ADMIN_PASSWORD"])

        UserCollection.delete_one(admin_login)

        admin = User(name=admin_name, surname=admin_surname, email=admin_email, id=admin_login,
                     password_hash=admin_password_hash)
        admin.is_admin = True
        UserCollection.save_data(admin)

    @staticmethod
    @login.user_loader
    def find_one(user_id: str) -> User:
        data = pymongo.db.users.find_one({"_id": user_id})
        if data is not None:
            return User.create_from_dict(data)

    @staticmethod
    def has_email(email: str) -> bool:
        return pymongo.db.users.find_one({"email": email}) is not None

    @staticmethod
    def find_one_or_404(user_id: str) -> User:
        data = pymongo.db.users.find_one_or_404({"_id": user_id})
        return User.create_from_dict(data)

    @staticmethod
    def find_all() -> List[User]:
        return [User.create_from_dict(data) for data in pymongo.db.users.find()]

    @staticmethod
    def save_data(data: User) -> None:
        data_dict = asdict(data)
        del data_dict["id"]
        pymongo.db.users.update_one({"_id": data.id}, {'$set': data_dict}, upsert=True)

    @staticmethod
    def delete_one(user_id: str) -> None:
        pymongo.db.users.delete_one({"_id": user_id})

    @staticmethod
    def docs_count() -> int:
        return pymongo.db.users.estimated_document_count()


class PatientCollection:

    @staticmethod
    def find_one(patient_id: str, cls: type = None) -> _PatientData:
        PatientCollection.__cls_check(cls)

        _cls = cls or Patient

        if _cls == Patient:
            data = pymongo.db.patients.find_one_or_404({"_id": ObjectId(patient_id)})
        else:
            data = pymongo.db.patients.find_one_or_404({"_id": ObjectId(patient_id)}, {cls.FIELD_NAME: 1})

        return _cls.create_from_dict(data)

    @staticmethod
    def find_all(cls: type = None) -> List[_PatientData]:
        _cls = cls or Patient
        return [_cls.create_from_dict(data) for data in pymongo.db.patients.find()]

    @staticmethod
    def delete_field(patient_id: str, cls: type) -> None:
        PatientCollection.__cls_check(cls)
        pymongo.db.patients.update_one({"_id": ObjectId(patient_id)}, {"$unset": {cls.FIELD_NAME: ""}})

    @staticmethod
    def save_data(data: _PatientData, patient_id: str = None) -> str:
        if patient_id:
            pymongo.db.patients.update_one({"_id": ObjectId(patient_id)}, {'$set': data.serialize()})
            return patient_id
        else:
            inserted = pymongo.db.patients.insert_one(data.serialize())
            return str(inserted.inserted_id)

    @staticmethod
    def paginate(page_num: int, cls: type = None, sort_rules: Dict[str, int] = None) -> Tuple[
        List[_PatientData], bool, bool]:
        PatientCollection.__cls_check(cls)

        _cls = cls or Patient

        if _cls == Patient:
            cursor = pymongo.db.patients.find()
        else:
            cursor = pymongo.db.patients.find({}, {_cls.FIELD_NAME: 1})

        if sort_rules is not None:
            cursor = cursor.sort(list(sort_rules.items()))

        page_size = current_app.config["PATIENTS_PAGE_SIZE"]
        skip_cnt = page_size * (page_num - 1)
        cursor = cursor.skip(skip_cnt).limit(page_size)

        docs = [_cls.create_from_dict(data) for data in cursor]

        has_next = pymongo.db.patients.estimated_document_count() > page_num * page_size
        has_prev = page_num > 1

        return docs, has_next, has_prev

    @staticmethod
    def __cls_check(cls: type = None) -> None:
        if cls is not None and not issubclass(cls, _PatientData):
            raise ValueError(f"Passed class must be a subclass of {_PatientData.__name__}")
